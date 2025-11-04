from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any
import json


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_branding(templates_dir: Optional[Path], branding: Optional[str]) -> Dict[str, Any]:
    if not templates_dir:
        return {}
    candidates = []
    if branding:
        candidates.append(templates_dir / f"{branding}.json")
    candidates.append(templates_dir / "branding.json")
    for p in candidates:
        if p.exists():
            try:
                return json.loads(p.read_text(encoding="utf-8", errors="ignore"))
            except Exception:
                return {}
    return {}


from models import DocFormatResult


class DocFormatterAgent:
    def __init__(self, templates_dir: Optional[Path] = None, output_dir: Optional[Path] = None) -> None:
        self.templates_dir = templates_dir
        self.output_dir = output_dir

    def format(self, input_path: Path, fmt: str = "md", branding: Optional[str] = None) -> DocFormatResult:
        text = _read_text(input_path)
        branding_data = _load_branding(self.templates_dir, branding)

        header = self._compose_header(branding_data)
        content_md = header + "\n\n" + text.strip() + "\n"

        out_dir = self.output_dir or input_path.parent
        out_dir.mkdir(parents=True, exist_ok=True)

        stem = input_path.stem
        if fmt == "md":
            out = out_dir / f"{stem}_formatted.md"
            out.write_text(content_md, encoding="utf-8")
            return DocFormatResult(input_path=input_path, output_path=out, requested_format=fmt, actual_format="md")

        if fmt == "docx":
            try:
                from docx import Document  # type: ignore
                from docx.shared import Pt, Inches  # type: ignore

                doc = Document()
                self._docx_apply_branding(doc, branding_data)
                # Optional logo image in templates dir
                logo_name = branding_data.get("logo")
                if logo_name and self.templates_dir:
                    logo_path = (self.templates_dir / logo_name)
                    if logo_path.exists():
                        try:
                            doc.add_picture(str(logo_path), width=Inches(1.8))
                        except Exception:
                            pass
                self._docx_add_markdownish(doc, text)
                out = out_dir / f"{stem}.docx"
                doc.save(out)
                return DocFormatResult(input_path=input_path, output_path=out, requested_format=fmt, actual_format="docx")
            except Exception:
                # Fallback to Markdown
                out = out_dir / f"{stem}_formatted.md"
                out.write_text(content_md, encoding="utf-8")
                return DocFormatResult(input_path=input_path, output_path=out, requested_format=fmt, actual_format="md")

        if fmt == "pdf":
            try:
                # Minimal PDF rendering
                from reportlab.lib.pagesizes import LETTER  # type: ignore
                from reportlab.pdfgen import canvas  # type: ignore
                from reportlab.lib.units import inch  # type: ignore

                out = out_dir / f"{stem}.pdf"
                c = canvas.Canvas(str(out), pagesize=LETTER)
                width, height = LETTER

                # Title
                title = branding_data.get("title", "Document")
                c.setFont("Helvetica-Bold", 16)
                c.drawString(1 * inch, height - 1 * inch, title)

                # Subtitle and date
                y = height - 1.3 * inch
                c.setFont("Helvetica", 10)
                subtitle = branding_data.get("subtitle", "")
                if subtitle:
                    c.drawString(1 * inch, y, subtitle)
                    y -= 0.2 * inch
                c.drawString(1 * inch, y, datetime.now().strftime("%Y-%m-%d"))
                y -= 0.4 * inch

                # Body text naive wrap
                c.setFont("Times-Roman", 11)
                max_width = width - 2 * inch
                for line in text.splitlines():
                    for chunk in self._wrap_line(line, c, max_width):
                        if y < 1 * inch:
                            c.showPage()
                            y = height - 1 * inch
                            c.setFont("Times-Roman", 11)
                        c.drawString(1 * inch, y, chunk)
                        y -= 0.18 * inch
                c.save()
                return DocFormatResult(input_path=input_path, output_path=out, requested_format=fmt, actual_format="pdf")
            except Exception:
                out = out_dir / f"{stem}_formatted.md"
                out.write_text(content_md, encoding="utf-8")
                return DocFormatResult(input_path=input_path, output_path=out, requested_format=fmt, actual_format="md")

        # Default: write markdown
        out = out_dir / f"{stem}_formatted.md"
        out.write_text(content_md, encoding="utf-8")
        return DocFormatResult(input_path=input_path, output_path=out, requested_format=fmt, actual_format="md")

    # ---------------------------- helpers -----------------------------
    def _compose_header(self, branding: Dict[str, Any]) -> str:
        title = branding.get("title", "Report")
        subtitle = branding.get("subtitle", "")
        author = branding.get("author", "")
        date_str = datetime.now().strftime("%Y-%m-%d")
        lines = [f"# {title}"]
        if subtitle:
            lines.append(f"_{subtitle}_")
        meta = [f"Date: {date_str}"]
        if author:
            meta.append(f"Author: {author}")
        lines.append(" | ".join(meta))
        return "\n".join(lines)

    def _docx_apply_branding(self, doc, branding: Dict[str, Any]) -> None:
        title = branding.get("title", "Report")
        subtitle = branding.get("subtitle")
        author = branding.get("author")
        date_str = datetime.now().strftime("%Y-%m-%d")

        h = doc.add_heading(title, 0)
        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.style = doc.styles["Intense Quote"] if "Intense Quote" in doc.styles else None
        meta = doc.add_paragraph(f"Date: {date_str}" + (f" | Author: {author}" if author else ""))
        doc.add_paragraph("")

    def _docx_add_markdownish(self, doc, text: str) -> None:
        # Minimal Markdown handling: headings, bullets, fenced code, inline bold/italic/code, links
        in_code = False
        code_lines = []
        for raw in text.splitlines():
            s = raw.rstrip("\n")
            if s.strip().startswith("```"):
                if in_code:
                    # flush code block
                    p = doc.add_paragraph("\n".join(code_lines))
                    try:
                        p.style = doc.styles.get("Code") or p.style
                    except Exception:
                        pass
                    # set monospace
                    for run in p.runs:
                        try:
                            run.font.name = "Courier New"
                            run.font.size = Pt(10)
                        except Exception:
                            pass
                    code_lines = []
                    in_code = False
                else:
                    in_code = True
                continue
            if in_code:
                code_lines.append(s)
                continue

            s = s.strip()
            if not s:
                doc.add_paragraph("")
                continue
            if s.startswith("### "):
                doc.add_heading(s[4:].strip(), level=3)
                continue
            if s.startswith("## "):
                doc.add_heading(s[3:].strip(), level=2)
                continue
            if s.startswith("# "):
                doc.add_heading(s[2:].strip(), level=1)
                continue
            if s.startswith("- "):
                self._docx_add_inline(doc.add_paragraph(style="List Bullet"), s[2:].strip())
                continue
            self._docx_add_inline(doc.add_paragraph(), s)

    def _docx_add_inline(self, paragraph, text: str) -> None:
        # Parse inline: **bold**, *italic*, `code`, [text](url)
        import re as _re
        tokens = []
        i = 0
        pattern = _re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))")
        for m in pattern.finditer(text):
            if m.start() > i:
                tokens.append(("text", text[i:m.start()]))
            tokens.append(("mark", m.group(0)))
            i = m.end()
        if i < len(text):
            tokens.append(("text", text[i:]))

        run = None
        for kind, val in tokens:
            if kind == "text":
                run = paragraph.add_run(val)
            else:
                if val.startswith("**") and val.endswith("**"):
                    r = paragraph.add_run(val[2:-2])
                    r.bold = True
                elif val.startswith("*") and val.endswith("*"):
                    r = paragraph.add_run(val[1:-1])
                    r.italic = True
                elif val.startswith("`") and val.endswith("`"):
                    r = paragraph.add_run(val[1:-1])
                    try:
                        r.font.name = "Courier New"
                    except Exception:
                        pass
                elif val.startswith("["):
                    # link: [text](url) -> "text (url)"
                    try:
                        t = val[1:val.index("]")]
                        u = val[val.index("(")+1:-1]
                        paragraph.add_run(f"{t} ({u})")
                    except Exception:
                        paragraph.add_run(val)

    def _wrap_line(self, line: str, canvas, max_width: float):
        # Break a line into chunks that fit the given width
        words = line.split()
        if not words:
            return [""]
        chunks = []
        cur = words[0]
        for w in words[1:]:
            test = cur + " " + w
            if canvas.stringWidth(test, "Times-Roman", 11) <= max_width:
                cur = test
            else:
                chunks.append(cur)
                cur = w
        chunks.append(cur)
        return chunks
